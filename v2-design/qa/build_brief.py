from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DESIGN = ROOT / "v2-design"
CONCEPTS = DESIGN / "concepts"
QA_BROWSER = DESIGN / "qa" / "browser"
OUT = DESIGN / "brief" / "Golden-Hour-V2-Product-Visual-Design-Brief.docx"

INK = "171A1C"
GRAPHITE = "101416"
AMBER = "A85F08"
AMBER_BRIGHT = "E9A23B"
IVORY = "F5F0E6"
MIST = "EBE4D7"
MUTED = "686963"
LINE = "D7CFBF"
RED = "B43F43"
GREEN = "276A51"
WHITE = "FFFAF1"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_hyperlink(paragraph, text, url, color=AMBER, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u_node = OxmlElement("w:u")
        u_node.set(qn("w:val"), "single")
        r_pr.append(u_node)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_alt_text(inline_shape, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description[:80])


def add_rule(paragraph, color=LINE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_text(doc, text, style=None, size=None, color=INK, bold=False, italic=False, before=0, after=6, align=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_kicker(doc, text, after=5):
    paragraph = add_text(doc, text.upper(), size=8.5, color=AMBER, bold=True, after=after)
    for run in paragraph.runs:
        run.font.letter_spacing = Pt(1.3) if hasattr(run.font, "letter_spacing") else None
    return paragraph


def add_callout(doc, label, title, body, fill=IVORY, accent=AMBER):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label.upper())
    set_run_font(run, size=8, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(title)
    set_run_font(run, size=15, color=INK, bold=True)
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    run = p3.add_run(body)
    set_run_font(run, size=10.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_numbered(doc, items: Iterable[tuple[str, str]]):
    for title, body in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        r1 = p.add_run(f"{title}. ")
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(body)
        set_run_font(r2, size=11, color=INK)


def add_figure(doc, image_path: Path, caption: str, alt: str, max_width=6.5, max_height=7.4, page_break=False):
    if page_break:
        doc.add_page_break()
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width = min(max_width, max_height * ratio)
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    add_alt_text(shape, alt)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_two_figures(
    doc,
    first: Path,
    second: Path,
    caption: str,
    alts: tuple[str, str],
    max_height=6.9,
    max_item_width=2.72,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    for index, (image_path, alt) in enumerate(((first, alts[0]), (second, alts[1]))):
        with Image.open(image_path) as image:
            w, h = image.size
        ratio = w / h
        height = max_height
        width = height * ratio
        # Keep both figures and their inter-image gap on the same line. Word's
        # layout engine otherwise wraps the second portrait image to a new page,
        # leaving an orphaned section title behind.
        if width > max_item_width:
            width = max_item_width
            height = width / ratio
        if index:
            p.add_run("     ")
        shape = p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
        add_alt_text(shape, alt)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, AMBER, 16, 8),
        "Heading 2": (13, AMBER, 12, 6),
        "Heading 3": (12, GRAPHITE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)

    if "Figure Page Title" not in styles:
        fig = styles.add_style("Figure Page Title", WD_STYLE_TYPE.PARAGRAPH)
        fig.font.name = "Calibri"
        fig.font.size = Pt(17)
        fig.font.bold = True
        fig.font.color.rgb = rgb(INK)
        fig.paragraph_format.space_before = Pt(0)
        fig.paragraph_format.space_after = Pt(8)
        fig.paragraph_format.keep_with_next = True


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("GOLDEN HOUR V2  |  PRODUCT + VISUAL DESIGN BRIEF")
        set_run_font(run, size=7.5, color=MUTED, bold=True)

        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        set_table_geometry(table, [7020, 2340], indent_dxa=0)
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = table.cell(0, 0).paragraphs[0].add_run("Independent hackathon prototype  |  25 August 2026")
        set_run_font(r, size=7.5, color=MUTED)
        p2 = table.cell(0, 1).paragraphs[0]
        r2 = p2.add_run("PAGE ")
        set_run_font(r2, size=7.5, color=MUTED, bold=True)
        add_page_field(p2)


def capability_table(doc):
    rows = [
        ("Four-question triage", "Preserved", "/action-room/intake", "Amount, route, time, authorisation"),
        ("Urgent action sequence", "Preserved", "/action-room/case", "1930, freeze, bank, evidence, FIR"),
        ("Recovery clock + layer model", "Preserved + relabeled", "/action-room/money-trail", "Illustrative model, not personal probability"),
        ("English/Hindi call script", "Preserved", "/action-room/case", "Bilingual 1930 guidance"),
        ("OpenAI chronology", "Preserved + verified", "/action-room/complaint", "Server key plus deterministic fallback"),
        ("Evidence checklist", "Preserved + promoted", "/action-room/evidence", "Dedicated page"),
        ("Recovery-stage machine", "Preserved + promoted", "/action-room/recovery", "Dedicated page"),
        ("Frozen-account test", "Preserved", "/action-room/frozen", "Position, bank letter, RTI"),
        ("Real vs. mocked", "Updated", "/methodology", "Current primary sources"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2300, 1550, 2200, 3310])
    repeat_table_header(table.rows[0])
    headers = ["Capability", "Status", "V2 route", "Notes"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, GRAPHITE)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=9, color=WHITE, bold=True)
    for capability, status, route, note in rows:
        cells = table.add_row().cells
        values = [capability, status, route, note]
        for index, value in enumerate(values):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, size=8.5, color=INK, bold=index == 0)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if status.startswith("Preserved"):
            set_cell_shading(cells[1], "E1EEE7")
        else:
            set_cell_shading(cells[1], "F3DFBD")
    set_table_geometry(table, [2300, 1550, 2200, 3310])


def route_table(doc):
    rows = [
        ("Public", "/", "Explain the product and route the situation"),
        ("Overview", "/action-room", "Resume or choose a workspace"),
        ("Triage", "/action-room/intake", "Create the local case"),
        ("Urgent actions", "/action-room/what-to-do", "Prioritized first-hour guidance"),
        ("Case", "/action-room/case", "Run and log the original action sequence"),
        ("Money trail", "/action-room/money-trail", "Visual layer model and assumptions"),
        ("Evidence", "/action-room/evidence", "Preservation checklist"),
        ("Complaint", "/action-room/complaint", "OpenAI chronology + bank notice"),
        ("Recovery", "/action-room/recovery", "Local actions and simulated stages"),
        ("Frozen", "/action-room/frozen", "Separate account-restraint workflow"),
        ("Methodology", "/methodology", "Sources, privacy, real vs. simulated"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 3000, 4560])
    repeat_table_header(table.rows[0])
    for index, text in enumerate(("Surface", "Route", "Single job")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, GRAPHITE)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=9, color=WHITE, bold=True)
    for surface, route, job in rows:
        cells = table.add_row().cells
        for index, value in enumerate((surface, route, job)):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, size=8.8, color=INK, bold=index == 0)
        if surface == "Public":
            set_cell_shading(cells[0], IVORY)
        elif surface == "Methodology":
            set_cell_shading(cells[0], MIST)
        else:
            set_cell_shading(cells[0], "EFE7D7")
    set_table_geometry(table, [1800, 3000, 4560])


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Cover: editorial_cover pattern with a brand-display override.
    add_text(doc, "GOLDEN HOUR", size=9, color=AMBER, bold=True, after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_text(doc, "Product + Visual Design Brief", size=34, color=GRAPHITE, bold=False, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.runs[0], name="Georgia", size=34, color=GRAPHITE)
    add_text(doc, "Premium public entry point → separate multi-page Cyber Fraud Action Room", size=14, color=MUTED, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Architecture reset • implemented V2 • visual approval and handoff", size=9.5, color=AMBER, bold=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure(
        doc,
        CONCEPTS / "01-visual-north-star.png",
        "North-star material language: warm evidence surface, interrupted transaction path, focused graphite command surface.",
        "Abstract Golden Hour visual direction with an ivory clock, evidence paper, interrupted path, and graphite containment graphics.",
        max_width=6.2,
        max_height=4.1,
    )
    add_text(doc, "Prepared for Aariz • 25 August 2026", size=10, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Independent OpenAI hackathon prototype • no government or bank affiliation", size=8.5, color=MUTED, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_kicker(doc, "Decision summary")
    doc.add_heading("The product is now an application, not a long landing page", level=1)
    add_callout(
        doc,
        "Architecture decision",
        "Short public entry. Deep operational workspace.",
        "The landing page communicates the promise and routes the situation. Every response job lives on a fresh page inside the Action Room, so urgent guidance never competes with marketing content.",
        fill=IVORY,
    )
    add_text(doc, "What changed", style="Heading 2")
    add_bullets(doc, [
        "The earlier eight-section landing-led concept is superseded as the primary information architecture.",
        "The public page now has four purposeful bands: emergency utility, product hero, situation routing, and a concise product/prototype boundary.",
        "The Action Room is a separate graphite application shell with dedicated pages for triage, urgent actions, money trail, evidence, complaint drafting, recovery tracking, and frozen accounts.",
        "All original working capabilities were ported into the V2 folder and verified end to end.",
        "The original repository root remains unchanged; V2 is isolated under /v2.",
    ])
    add_text(doc, "Approval position", style="Heading 2")
    add_text(doc, "This brief documents the implemented direction. The core architecture, visual system, and functionality are working; remaining work is product polish, legal review, and deployment choice rather than rebuilding the concept.", size=11, color=INK, after=8)

    doc.add_page_break()
    add_kicker(doc, "Information architecture")
    doc.add_heading("One job per page", level=1)
    add_text(doc, "The route map is deliberately explicit. A hackathon judge can understand the product quickly, while a distressed user can enter a task without scanning a marketing page.", size=11, color=INK, after=10)
    route_table(doc)
    add_text(doc, "Compatibility", style="Heading 2")
    add_text(doc, "Legacy routes /act, /case, /frozen, and /how-it-works redirect to their V2 equivalents. Existing demo links therefore continue to work.", size=10.5, color=MUTED)

    doc.add_page_break()
    add_kicker(doc, "Capability parity")
    doc.add_heading("Original functionality, reorganized and preserved", level=1)
    add_text(doc, "The architecture changed; the product capabilities did not disappear. This matrix is the implementation contract used for the end-to-end test.", size=11, after=10)
    capability_table(doc)

    doc.add_page_break()
    add_kicker(doc, "Experience model")
    doc.add_heading("Two surfaces with different jobs", level=1)
    add_text(doc, "Public surface", style="Heading 2")
    add_text(doc, "Warm ivory, editorial typography, calm explanation, one generated hero asset, and strong situation routing. This surface earns trust and gets out of the way.", size=11)
    add_text(doc, "Action Room surface", style="Heading 2")
    add_text(doc, "Deep graphite, amber task focus, persistent emergency access, local-case status, and dedicated navigation. This surface is a tool, not a campaign.", size=11)
    add_text(doc, "Victim path", style="Heading 2")
    add_numbered(doc, [
        ("Triage", "Answer amount, payment route, time window, and authorisation."),
        ("Act", "Call 1930, prepare the freeze request, notify the bank, preserve evidence, and prepare FIR information."),
        ("Document", "Use evidence to build an editable complaint chronology and written bank notice."),
        ("Track", "Log actions, references, simulated stages, and next follow-up."),
    ])
    add_text(doc, "Frozen-account path", style="Heading 2")
    add_numbered(doc, [
        ("Diagnose", "Capture state, duration, balance held, disputed amount, legal section, scope, notice, and Magistrate information."),
        ("Understand", "Show the strength and limits of the position using sourced rules."),
        ("Respond", "Generate a bank letter and RTI draft without pretending to provide legal representation."),
        ("Escalate", "Preserve the written record and track review points."),
    ])

    doc.add_page_break()
    add_kicker(doc, "Visual system")
    doc.add_heading("Calm emergency intelligence", level=1)
    add_callout(doc, "Palette", "Ivory → graphite, with amber as the operational signal", "Public: Ivory #F5F0E6, Ink #171A1C, warm paper grain. App: Graphite #101416, warm white #F5F0E6, Amber #E9A23B. Red is reserved for destructive or overdue states.", fill=IVORY)
    add_text(doc, "Typography", style="Heading 2")
    add_bullets(doc, [
        "DM Serif Display carries public headlines and app page titles as an editorial signal.",
        "Manrope carries controls, navigation, forms, and operational copy.",
        "JetBrains Mono carries case IDs, timing, stage metadata, and numeric values.",
        "Noto Sans Devanagari supports Hindi guidance without substituting person imagery or flags.",
    ])
    add_text(doc, "Component language", style="Heading 2")
    add_bullets(doc, [
        "Rounded task cards with precise borders and shallow depth; selective glass only where focus or hierarchy benefits.",
        "Thin Phosphor icons for clocks, paths, documents, banks, phones, calendars, exports, and evidence. No shield/padlock identity and no person icon.",
        "Progress rails, audit timelines, transaction paths, status chips, evidence checklists, and generated-document panes replace long explanatory text.",
        "Motion is limited to transform and opacity, supports reduced motion, and never delays an urgent action.",
    ])

    doc.add_page_break()
    add_text(doc, "Implemented public landing", style="Figure Page Title")
    add_figure(doc, QA_BROWSER / "desktop-landing.png", "Live V2 desktop landing capture.", "Full desktop screenshot of the warm ivory Golden Hour public landing page.", max_width=4.15, max_height=7.65)

    doc.add_page_break()
    add_text(doc, "Implemented Action Room", style="Figure Page Title")
    add_figure(doc, QA_BROWSER / "desktop-action-room.png", "Live V2 desktop Action Room overview capture.", "Desktop screenshot of the deep graphite Action Room with persistent navigation and workspace cards.", max_width=6.5, max_height=6.8)
    add_text(doc, "The public/app transition is not a theme toggle. It is a deliberate context switch from explanation to operation.", size=10, color=MUTED, italic=True)

    doc.add_page_break()
    add_text(doc, "Implemented case workflow", style="Figure Page Title")
    add_figure(doc, QA_BROWSER / "desktop-case.png", "Live V2 case screen after triage.", "Desktop screenshot of the Golden Hour case workspace with recovery clock, money layer map, call script, and urgent actions.", max_width=6.5, max_height=7.15)

    doc.add_page_break()
    add_text(doc, "Implemented mobile surfaces", style="Figure Page Title")
    add_two_figures(
        doc,
        QA_BROWSER / "mobile-landing.png",
        QA_BROWSER / "mobile-action-room.png",
        "Live mobile captures: public landing (left) and separate Action Room application (right).",
        (
            "Tall mobile screenshot of the Golden Hour public landing page.",
            "Tall mobile screenshot of the Golden Hour Action Room with bottom navigation.",
        ),
        max_height=7.25,
    )

    doc.add_page_break()
    add_kicker(doc, "Generated visual direction")
    doc.add_heading("North star and operational concepts", level=1)
    add_text(doc, "The image-generation work established the material language before implementation: clocks, interrupted transaction paths, evidence documents, bank/process objects, and graphite containment surfaces. It deliberately excludes people, faces, hands, silhouettes, avatars, shields, and padlocks.", size=11, after=8)
    add_figure(doc, CONCEPTS / "12-desktop-action-room-overview.png", "Generated desktop Action Room concept used to set hierarchy and task density.", "Generated deep graphite desktop Action Room interface concept with amber task focus and no people.", max_width=6.5, max_height=4.5)

    doc.add_page_break()
    add_text(doc, "Victim flow concepts", style="Figure Page Title")
    add_two_figures(doc, CONCEPTS / "13-mobile-victim-triage.png", CONCEPTS / "14-mobile-victim-act-now.png", "Generated victim path: triage and act-now stages.", ("Generated mobile triage screen.", "Generated mobile urgent-action screen."), max_height=6.8)

    doc.add_page_break()
    add_text(doc, "Victim flow concepts — continued", style="Figure Page Title")
    add_two_figures(doc, CONCEPTS / "15-mobile-victim-build-evidence.png", CONCEPTS / "16-mobile-victim-track-recovery.png", "Generated victim path: evidence and recovery tracking stages.", ("Generated mobile evidence-building screen.", "Generated mobile recovery-tracking screen."), max_height=6.8)

    doc.add_page_break()
    add_text(doc, "Frozen-account concepts", style="Figure Page Title")
    add_two_figures(doc, CONCEPTS / "17-mobile-frozen-diagnose.png", CONCEPTS / "18-mobile-frozen-understand.png", "Generated frozen-account path: diagnose and understand stages.", ("Generated mobile frozen-account diagnosis screen.", "Generated mobile frozen-account position screen."), max_height=6.8)

    doc.add_page_break()
    add_text(doc, "Frozen-account concepts — continued", style="Figure Page Title")
    add_two_figures(doc, CONCEPTS / "19-mobile-frozen-generate-response.png", CONCEPTS / "20-mobile-frozen-track-escalation.png", "Generated frozen-account path: response drafting and escalation tracking.", ("Generated mobile frozen-account response screen.", "Generated mobile frozen-account escalation screen."), max_height=6.8)

    doc.add_page_break()
    add_kicker(doc, "Model, privacy, and safety")
    doc.add_heading("The model is bounded; the emergency workflow is deterministic", level=1)
    add_text(doc, "OpenAI use", style="Heading 2")
    add_bullets(doc, [
        "Only the complaint chronology uses the OpenAI API.",
        "The key is read server-side from OPENAI_API_KEY or the existing local OPENAI_KEY variable; it is never exposed through NEXT_PUBLIC_.",
        "The system prompt forbids invented identifiers, preserves the authorisation branch, and requires missing-detail placeholders.",
        "A deterministic template completes the workflow if the model is unavailable, rejected, or returns empty content.",
        "The production smoke test returned a real model-generated draft during verification.",
    ])
    add_text(doc, "Local data", style="Heading 2")
    add_bullets(doc, [
        "No sign-up, database, or case analytics.",
        "The case record is versioned in browser local storage and can be erased from the Action Room.",
        "Only an explicit drafting action sends the narrative and minimal incident context to the server.",
        "The UI repeatedly warns against Aadhaar, PAN, OTP, PIN, passwords, and full card details.",
    ])
    add_text(doc, "Public-safety language", style="Heading 2")
    add_bullets(doc, [
        "The product says REPORT NOW and explains why speed supports containment; it does not publish a personal recovery probability.",
        "1930 and cybercrime.gov.in are presented as official channels.",
        "Frozen-account outputs say general information, not legal advice, and distinguish sourced rules from case-specific review.",
        "All synthetic references and simulated stages are labeled at the point of use.",
    ])

    doc.add_page_break()
    add_kicker(doc, "Primary sources")
    doc.add_heading("Current procedural and factual references", level=1)
    sources = [
        ("Ministry of Home Affairs parliamentary answer, 24 March 2026", "https://www.mha.gov.in/MHA1/Par2017/pdfs/par2026-pdfs/LS24032026/5124.pdf", "Official reporting workflow and national channels."),
        ("Press Information Bureau release, 21 August 2026", "https://www.pib.gov.in/PressReleasePage.aspx?PRID=2301973&lang=2&reg=48", "As of 30 June 2026, Rs 11,158 crore saved across 32.80 lakh complaints; an aggregate, not a personal probability."),
        ("RBI customer-protection circular RBI/2017-18/15", "https://www.rbi.org.in/commonman/English/scripts/Notification.aspx?Id=2623", "Unauthorised electronic-banking liability and notification timelines."),
        ("Bharatiya Nagarik Suraksha Sanhita, 2023", "https://www.indiacode.nic.in/bitstream/123456789/21920/1/the_bharatiya_nagarik_suraksha_sanhita%2C_2023.pdf", "Official text for sections 94, 106, and 107."),
        ("Delhi High Court judgment, 2026", "https://delhihighcourt.nic.in/app/showlogo/1769248155_eb87a6191c1f9386_596_41982025.pdf/2026", "Discussion of debit-freeze power under section 106 and the separate section 107 process."),
    ]
    for title, url, note in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        set_run_font(r, size=10.5, color=INK, bold=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.2)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(note)
        set_run_font(r2, size=9.5, color=MUTED)
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.2)
        p3.paragraph_format.space_after = Pt(9)
        add_hyperlink(p3, url, url, color=AMBER)

    doc.add_page_break()
    add_kicker(doc, "Verification and handoff")
    doc.add_heading("Implementation acceptance checklist", level=1)
    checklist = [
        "Production build passes across 19 routes.",
        "ESLint passes with zero warnings.",
        "Desktop and mobile landing surfaces render without overlays or missing assets.",
        "Action Room navigation works across dedicated pages.",
        "Four-question triage creates and persists a local case.",
        "English and Hindi 1930 call guidance works.",
        "Evidence selection and urgent-action logging work.",
        "OpenAI draft generation works with the server-side key; fallback remains available.",
        "Recovery tracker responds to local actions.",
        "Frozen-account diagnostic produces position, bank letter, and RTI output.",
        "Legacy route redirects work.",
        "No people, faces, hands, silhouettes, avatars, or person icons appear in project visuals.",
        "Original root prototype files remain untouched.",
    ]
    add_bullets(doc, [f"PASS — {item}" for item in checklist])
    add_callout(doc, "Handoff", "The architecture reset is complete", "The next product decisions are deployment target, legal review ownership, and whether the full interface should receive a complete Hindi localization beyond the preserved bilingual emergency script.", fill=IVORY)

    # Author metadata belongs to Aariz, never Codex.
    doc.core_properties.title = "Golden Hour V2 Product + Visual Design Brief"
    doc.core_properties.subject = "Multi-page Cyber Fraud Action Room architecture, visual system, and implementation handoff"
    doc.core_properties.author = "Aariz"
    doc.core_properties.last_modified_by = "Aariz"
    doc.core_properties.keywords = "Golden Hour, cyber fraud, action room, OpenAI hackathon, product design"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
